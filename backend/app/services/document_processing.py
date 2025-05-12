from flask import current_app
import PyPDF2
import docx
import spacy
from collections import defaultdict
from tenacity import retry, stop_after_attempt, wait_exponential
from app.config import Config
from app.utils.nlp_utils import RISK_KEYWORDS, AMBIGUOUS_TERMS
from gcp.gcp_client import extract_clauses, analyze_entities, vertex_summarize
from google.cloud import language_v1  # Added for Document and Entity Type

# Load SpaCy model once
_nlp = spacy.load("en_core_web_sm")

def extract_text_from_pdf(file_path):
    try:
        with open(file_path, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            return ''.join(p.extract_text() or "" for p in reader.pages)
    except Exception as e:
        current_app.logger.error(f"PDF extraction error: {e}")
        return None

def extract_text_from_docx(file_path):
    try:
        doc = docx.Document(file_path)
        return "\n".join(p.text for p in doc.paragraphs)
    except Exception as e:
        current_app.logger.error(f"DOCX extraction error: {e}")
        return None

def preprocess_text(text):
    text = " ".join(text.split())
    return "".join(c for c in text if c.isalnum() or c in [" ", ".", ",", "\n"])

def extract_legal_entities(text):
    try:
        # Create a Document object for the Natural Language API
        document = language_v1.Document(
            content=text,
            type_=language_v1.Document.Type.PLAIN_TEXT,
            language="en"
        )
        # Analyze entities using the NLP client from Config
        response = Config.NLP_CLIENT.analyze_entities(document=document)

        entities = defaultdict(list)
        for ent in response.entities:
            # Use the correct enum for entity type
            ent_type = language_v1.Entity.Type(ent.type_).name
            if ent_type in ["ORGANIZATION", "PERSON", "DATE", "LAW"]:
                entities[ent_type].append(ent.name)

        # Add clause detections using spaCy
        doc = _nlp(text)
        for sent in doc.sents:
            low = sent.text.lower()
            if "governing law" in low:
                entities.setdefault("CLAUSES", []).append("Governing Law")
            if "force majeure" in low:
                entities.setdefault("CLAUSES", []).append("Force Majeure")

        return dict(entities)
    except Exception as e:
        current_app.logger.error(f"Error analyzing entities: {e}")
        return {}  # Return empty dict on error

@retry(stop=stop_after_attempt(3), wait=wait_exponential(min=2, max=10))
def generate_summary(text):
    """Summarize via Vertex AI model loaded in Config."""
    try:
        response = Config.VERTEX_MODEL.predict(
            text,
            max_output_tokens=150,
            temperature=0.2
        )
        return response.text
    except Exception as e:
        current_app.logger.error(f"Vertex summarization error: {e}")
        return "Summary unavailable – API error"

def identify_legal_risks(text):
    risks = []
    doc = _nlp(text.lower())

    if not any("governing law" in s.text for s in doc.sents):
        risks.append("Missing Governing Law clause")

    risks += [f"Ambiguous term: {term}" for term in AMBIGUOUS_TERMS if term in text]

    for cat, kws in RISK_KEYWORDS.items():
        if any(kw in text.lower() for kw in kws):
            risks.append(f"Potential risk in {cat} clause")
    return risks

def extract_clauses_from_document(file_path):
    """Extract contract clauses using Document AI via gcp.gcp_client."""
    try:
        return extract_clauses(file_path)
    except Exception as e:
        current_app.logger.error(f"Clause extraction error: {e}")
        return []